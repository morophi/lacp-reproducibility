from pathlib import Path
import shutil

import docx
from docx.oxml.ns import qn


SRC = Path("experiment/lacp_ijibc_rev8.1.docx")
OUT = Path("experiment/lacp_ijibc_rev8.1_kor.docx")


KEY_TERMS = [
    "Retrieval-Augmented Generation (RAG)",
    "Large Language Models (LLMs)",
    "Local Agent Context Protocol (LACP)",
    "SC-Protocol (Stochastic Control Protocol)",
    "Context-Level Causal Intervention Model",
    "Information Augmentation Model",
    "Logit Margin Score (LMS)",
    "Modality Analysis (MA)",
    "Contextual Drift Score (CDS)",
    "Self-Reference Rate (SRR)",
    "Structural Consistency Index (SCI)",
    "Causal Intervention",
    "context-level intervention",
    "Counterfactual",
    "do-calculus",
    "Node A",
    "Node B",
    "Node C",
    "Test Run (TR)",
    "Calibration Run (CR)",
    "Calibration Run 2 (CR2)",
    "Run B",
    "CF-F",
    "LACP",
    "SC-Protocol",
    "RAG",
    "LLM",
    "LMS_delta",
    "MA_assert",
    "theta_MA",
    "theta_config.json",
]


PARA = {
    1: "**RAG**는 인과적으로 개입하는가? On-Premises **LLM** 기반 공공복지 상담 시스템에서 **Retrieval-Augmented Generation**의 **Causal Intervention Mechanism** 검증",
    3: "¹독립연구자, Korea (morophi@gmail.com)",
    4: "²한양사이버대학교 응용AI소프트웨어학과, Seoul, Korea",
    5: "원고 접수: xxxxxx / 수정: xxxxxx / 게재 승인: xxxxxx",
    6: "*교신저자: yah0612@hycu.ac.kr",
    7: "초록",
    8: "**Retrieval-Augmented Generation (RAG)**는 model weights를 수정하지 않고 **Large Language Models (LLMs)**에 외부 지식을 주입하는 방법으로 널리 채택되고 있다. 그러나 핵심 질문은 여전히 남아 있다. **RAG**는 단순히 정보를 보강하는가, 아니면 관찰 가능한 응답 구조를 변화시키는 externally controlled **context-level intervention**으로 작동할 수 있는가? 본 논문은 **RAG**가 **LLM**의 응답 방향과 confidence에 영향을 주는 controlled **context-level intervention**으로 operationalize될 수 있는지를 실증적으로 검토하기 위해 설계된 **Local Agent Context Protocol (LACP)** experiment를 제시한다. 우리는 on-premises 공공복지 행정 환경에서 세 개 노드의 simultaneous comparison architecture를 제안한다. 여기서 **Node A**(**RAG** + **SC-Protocol**), **Node B**(**RAG** only), **Node C**(Baseline)는 동일한 utterance를 동시에 수신한다. do-calculus-inspired intervention framework를 operationalize하여, 다섯 개 measurement dimensions인 **Logit Margin Score (LMS)**, **Modality Analysis (MA)**, **Contextual Drift Score (CDS)**, **Self-Reference Rate (SRR)**, **Structural Consistency Index (SCI)**를 정의하고, reverse causality problem을 직접 다루기 위해 **CF-F**(forced random intervention)를 포함한 여섯 개 **Counterfactual** conditions를 설계한다. 본 논문은 제안 연구의 experimental design, measurement methodology, infrastructure architecture를 제시한다. empirical results는 experimental runs 완료 후 보고될 예정이다.",
    9: "Keywords:",
    10: "**Retrieval-Augmented Generation**, **Causal Intervention**, **Large Language Model**, **On-Premises AI**, 공공복지 행정, **Logit Margin Score**, **Counterfactual Design**",
    11: "1. 서론",
    12: "**Large Language Models (LLMs)**는 natural language understanding and generation에서 뛰어난 역량을 보여 주며 다양한 domain에 도입되고 있다. 그러나 public sector environment에서 **LLM**을 배치하는 일은 고유한 과제를 동반한다. 지방정부 복지행정 시스템은 privacy and security regulations로 인해 엄격한 network isolation(air-gap) 제약 안에서 운영되어야 하며, 이는 cloud-based **LLM** services 사용을 사실상 금지한다. 또한 복지정책이 자주 개정되기 때문에 fine-tuning 기반 knowledge injection은 높은 유지보수 비용으로 인해 경제적으로 실용적이지 않다 [1].",
    13: "**Retrieval-Augmented Generation (RAG)**는 이러한 문제에 대한 유망한 해법으로 등장했다 [2]. **RAG**는 model weights를 수정하지 않고 document repositories의 외부 지식을 동적으로 제공함으로써, on-premises environment에서도 **LLM**이 최신 정보에 접근할 수 있게 한다. 그러나 근본적인 이론적 질문이 남아 있다. **RAG**는 단순한 information augmentation tool로 기능하는가, 아니면 judgment direction과 confidence structure에 측정 가능한 영향을 주는 externally controlled **context-level intervention**으로 작동하는가?",
    14: "이 구분은 public AI systems의 design and governance에 중요한 실천적 함의를 갖는다. **RAG**가 model이 무엇을 말하는지만이 아니라 얼마나 확신 있게, 어떤 구조로 말하는지를 바꾸는 **Causal Intervention**으로 기능한다면, high-stakes welfare consultation contexts에 이를 배치할 때 accountability와 explainability를 신중히 고려해야 한다. 본 연구는 **Local Agent Context Protocol (LACP)** experiment를 통해 이 공백을 다룬다.",
    15: "**LACP**는 model weights를 수정하지 않고 **RAG** intervention timing과 context injection method를 통제하는 experimental control layer로 operationalize된다. **SC-Protocol (Stochastic Control Protocol)** [3]은 **LACP** 내부의 trigger condition rule set이며, probabilistic AI systems가 human operators로부터 decision-making responsibility를 흡수해서는 안 된다는 human oversight principle에 기반한다 [4, 5, 6, 7]. 본 논문은 controlled experimental setting에서 **LACP**와 **SC-Protocol**을 처음 실증적으로 조사한다. **SC-Protocol**은 독립 construct로 평가되는 것이 아니라 intervention conditions를 조절하는 operational mechanism으로 기능한다.",
    16: "핵심 research question은 다음과 같다. '**RAG**는 simple information augmentation tool인가, 아니면 **LLM** judgment direction을 측정 가능하게 재구조화하는 externally controlled **context-level intervention**인가?' 우리는 Korean-language public welfare consultation scenario에서 three-node simultaneous comparison architecture를 통해 이 질문을 operationalize하고, fixed model weights, fixed user utterances, deterministic decoding, matched concurrent baseline conditions에 do-calculus-inspired **Causal Intervention** framework를 적용한다.",
    17: "2. 관련 연구",
    18: "2.1 Retrieval-Augmented Generation",
    19: "Lewis et al. [2]는 parametric memory와 non-parametric memory를 결합한 general-purpose framework로 **RAG**를 제안했다. 후속 연구는 **RAG**를 medical consultation [8], legal reasoning [9], enterprise knowledge management [10]로 확장했다. Shi et al. [11]은 irrelevant retrieved passages가 model outputs를 체계적으로 왜곡할 수 있음을 보였고, 이는 **RAG**의 영향이 단순한 information provision을 넘어섬을 시사한다. 기존 연구 중 do-calculus-inspired experimental framework에서 이러한 영향을 controlled **context-level intervention**으로 operationalize할 수 있는지를 실증적으로 검토한 연구는 없다. 또한 Graph RAG와 같은 최근 발전은 structured knowledge graphs와 text retrieval을 통합하여 global query-focused summarization을 개선했다 [12].",
    20: "2.2 LLM Confidence Measurement",
    21: "Kadavath et al. [13]은 pre-softmax logit values가 post-softmax probabilities보다 internal confidence states를 더 직접적으로 반영함을 보였다. Kuhn et al. [14]은 output uncertainty 측정치로 semantic entropy를 제안했다. Xiong et al. [15]은 **LLM** confidence elicitation methods를 survey하면서 verbalized confidence와 logit-based approaches를 구분했다.",
    22: "2.3 AI Systems에서의 Causal Inference",
    23: "Pearl의 **do-calculus** [16]는 observational conditioning과 controlled intervention을 구분하기 위한 theoretical reference point를 제공한다. Feder et al. [17]은 NLP systems에 causal analysis를 적용했다. 본 연구는 simultaneous controlled conditions 아래 do-calculus-inspired comparison framework를 operationalize함으로써 이 논의를 **RAG** systems로 확장한다.",
    24: "2.4 Public Sector의 On-Premises AI",
    25: "Wirtz et al. [18]은 public sector AI adoption의 주요 장벽으로 network isolation, data sovereignty, policy volatility를 제시했다. J. Kim and Lee [19]는 Korean local government contexts에서 **LLM** deployment challenges를 검토했다. 공공복지 상담을 위한 on-premises **LLM** systems에서 **RAG**가 controlled **context-level intervention**으로 기능할 수 있는지를 실증적으로 검토한 기존 연구는 없다. 엄격한 isolation 아래 on-premises systems를 구현하는 것은 regulatory requirement일 뿐 아니라 local governance에서 public trust와 technology acceptance를 확립하는 중요한 요인이다 [19, 20].",
    26: "3. 연구 설계",
    27: "3.1 Research Hypothesis",
    28: "H0 (**Information Augmentation Model**): **RAG**는 supplementary reference information을 제공하며, 관찰된 response changes는 concurrent baseline 대비 systematic intervention effects 없이 natural response variation, contextual accumulation, retrieval-related confidence increase로 설명될 수 있다.",
    29: "H1 (**Context-Level Causal Intervention Model**): **RAG**는 fixed model weights와 deterministic execution conditions 아래 concurrent baseline 대비 response direction, confidence structure, policy alignment에 statistically distinguishable changes를 산출하는 externally controlled **context-level intervention**으로 기능한다.",
    30: "3.2 LACP and SC-Protocol",
    31: "**Local Agent Context Protocol (LACP)**는 model weights를 수정하지 않고 external memory access management를 통해 **LLM** judgment intervention timing and method를 조절하는 on-premises agent control protocol이다. **SC-Protocol**은 **LACP** 내부의 trigger condition rule set이다. 본 연구에서 **SC-Protocol**은 measurement variable이 아니다. **Node A** conditions 전반에 대칭적으로 적용되므로 experimental treatment가 아니라 structural constraint로 기능한다. 따라서 이론적 비교는 세 개 observable contrasts로 분해된다. **Node B** minus **Node C**는 **RAG**-only context effect를 추정하고, **Node A** minus **Node B**는 intervention-governing constraint로서 **SC-Protocol**의 differential contribution을 추정하며, **Node A** minus **Node C**는 **RAG** plus **SC-Protocol**의 combined effect를 추정한다.",
    32: "3.3 Experimental Architecture",
    33: "본 실험은 three-node simultaneous comparison architecture를 사용한다. 세 개 **LLM** inference nodes는 동일한 utterances를 동시에 수신하고 독립적인 responses를 생성한다. 모든 nodes는 동일한 platform과 **LLM** model(Qwen series, on-premises Vulkan backend)을 사용한다.",
    34: "Table 1. Three-Node Architecture",
    35: "모든 nodes는 temperature=0.0(greedy decoding)으로 작동하여 stochastic sampling variance를 구조적으로 제거한다. residual inter-run variance는 **Calibration Run (CR)**(N=10)을 통해 정량화한다. N=10은 N=5 대비 standard error를 약 29% 감소시킨다(SE = σ/√N).",
    36: "본 실험에 사용된 model family는 explicit reasoning-mode outputs를 지원하므로, 모든 experimental runs는 reasoning/thinking output을 비활성화한 상태에서 실행된다. **CR** 이전에 logging pipeline은 생성된 response에 thinking-mode traces가 없는지 검증한다. 이 control은 reasoning-output artifacts가 **LMS**, **MA**, **CDS**, **SRR**, **SCI** measurements를 오염시키는 것을 방지한다.",
    37: "Node clocks는 internal chrony/NTP reference server를 통해 동기화된다. 본 실험은 simultaneity를 sub-millisecond physical execution simultaneity가 아니라 turn-dispatch level에서 다룬다. timestamp skew는 diagnostic purposes로 logging되며, causal comparison은 absolute wall-clock ordering이 아니라 within-turn matched responses에 기반한다.",
    38: "3.4 Experimental Run Structure",
    39: "Table 2. Experimental Run Structure",
    40: "**Test Run (TR)**은 세 nodes가 모두 활성화된 **RAG**-on conditions에서 수행되는 mandatory pre-CR gate이다. **TR**은 **Calibration Run (CR)**이 시작되기 전에 corpus canonicalization, chunking, embedding construction, ChromaDB ingestion, Harness orchestration, metric logging, MariaDB writes가 operational한지 검증한다. 이를 통해 preprocessing artifacts, routing failures, missing logs, unstable corpus conditions가 **CR** variance estimates를 오염시키는 것을 방지한다.",
    41: "**TR** acceptance는 corpus hash fixation, chunking and embedding audit, end-to-end pipeline integrity, **Node A**, **Node B**, **Node C**의 complete non-null outputs, **Node C** **RAG** contamination 없음, expected hardware/runtime jitter range 안의 node output differences를 요구한다. 실패한 **TR** runs는 infrastructure diagnosis를 위해 archive되며 causal effect estimation에는 사용되지 않는다.",
    42: "세부 **TR** substeps는 main causal design의 변경이 아니라 operational validity controls로 취급된다. 따라서 manuscript는 causal interpretability에 필요한 수준에서 validation logic을 보고하고, implementation-level checklists는 accompanying experiment guideline과 node checklist에 유지한다.",
    43: "top-k acceptance, hybrid retrieval exposure logging, response completeness verification, reasoning-output suppression, truncation monitoring을 포함한 추가 implementation-level controls는 accompanying experiment guideline and checklist에 정의된다. manuscript에서는 이러한 controls를 additional experimental treatments나 outcome variables가 아니라 pre-calibration eligibility criteria로 취급한다.",
    44: "formal retrieval candidate freeze gate도 **CR** 이전에 적용된다. 현재 candidate는 full-guideline corpus에서 body-only embedding surface를 사용해 구축된다. 반복되는 source-boundary와 provenance wrappers는 vectorized text에서 제외되지만, stored chunk document와 metadata는 audit 및 prompt display를 위해 provenance를 보존한다. 이 policy는 preliminary retrieval diagnostics에서 관찰된 source-boundary dominance를 다루며 retrieval substrate를 experimental outcome이 아니라 operational precondition으로 고정한다.",
    45: "선택된 pre-CR candidate는 vector top-30과 lexical top-30 candidate generation에 lexical-weighted reciprocal-rank fusion을 사용한다. source-aware table routing과 evidence sidecar generation은 diagnostic logging mechanisms로만 유지되며, **CR** 또는 **CR2**에서 default prompt-injection treatments로 사용되지 않는다. formal retrieval freeze는 stratified manual review가 section-, chunk-, block-type-, table-level evidence exposure를 검증할 때까지 보류된다. 이러한 controls는 **RAG**의 causal effect에 대한 empirical evidence가 아니라 retrieval-validity safeguards로 보고된다.",
    46: "3.5 Thermal Safeguard and Turn-Level Synchronization",
    47: "experimental environment는 동시에 작동하는 세 개 on-premises inference nodes를 사용한다. preliminary thermal-only direct inference probing은 짧은 five-turn load만으로도 node temperatures가 high-40 C range에서 약 69-74 C까지 상승할 수 있고, 10-second post-load observation이 maximum temperatures를 약 4.0-6.6 C 낮춘다는 것을 보였다. 관찰된 thermal behavior는 active fan cooling이 선택적 comfort setting이 아니라 inference runs를 stable operating envelope 안에 유지하기 위해 필요한 infrastructure control임을 보여준다.",
    48: "thermal drift가 uncontrolled confounder가 되는 것을 막기 위해 모든 experimental RUN stages는 turn-level synchronization barrier를 적용한다. 각 utterance에 대해 Harness는 동일한 turn index 아래 **Node A**, **Node B**, **Node C**로 requests를 dispatch하고, 세 node responses가 모두 완료되어 logging된 뒤에야 다음 turn으로 진행한다. 이 policy는 individual node completion times가 다르더라도 within-turn comparability를 보존한다.",
    49: "별도의 batch-level thermal safeguard는 completed turns 다섯 개마다 적용된다. batch boundary에서 Harness는 모든 inference nodes의 pre-cooldown temperature snapshot을 기록하고, 다음 dispatch를 30초 pause한 뒤 post-cooldown temperature snapshot을 기록하며, 이후 다음 turn batch를 재개한다. cooldown policy는 **TR**, **CR**, **CR2**, **CF**, **CF-F** runs 전반에 균일하게 적용되므로 experimental treatment가 아니라 fixed operational control로 취급된다.",
    50: "Table 2A. Rev. 8 Control Policy 이전 Thermal-only Direct Inference Probe",
    51: "thermal-only probe는 Agent, Harness-to-dblog run logging, **RAG** retrieval, validator execution, database writes를 의도적으로 우회했다. 유일한 persisted artifact는 local thermal JSONL file이었으며, 이를 통해 thermal characterization이 experimental result tables나 causal measurement logs를 오염시키지 않도록 했다.",
    52: "cooldown과 barrier rules는 모든 conditions에 대칭적으로 적용되므로 **RAG**-present와 **RAG**-absent conditions 사이의 causal contrast를 변경하지 않는다. 그 역할은 hardware-temperature-induced variance를 낮추고, long RUN stages 동안 response omission, node throttling, emergency shutdown의 probability를 낮춤으로써 data persistence reliability를 유지하는 것이다.",
    53: "Rev. 8.1 empirical update: follow-up thermal-only validation run은 parallel 1-second node sampling과 60-second cooldown observation window(run_id=thermal_only_5turn_cooldown60_parallel_20260526T170010)를 사용했다. 이 run은 15개 direct inference requests를 failure 없이 완료했다. 60-second cooldown interval 동안 node temperatures는 inference1에서 72.0 C에서 51.875 C로, inference2에서 69.0 C에서 52.25 C로, inference3에서 73.0 C에서 52.125 C로 감소했다.",
    54: "관찰된 cooldown profile은 simple linear cooling assumption으로 잘 설명되기보다 early-fast and late-slow pattern을 보인다. 따라서 thermal safeguard는 measured nonlinear post-load cooling behavior로 뒷받침되는 operational validity control로 문서화되며, throughput과 wall-clock duration은 primary causal outcome interpretation의 범위 밖에 둔다.",
    55: "4. Measurement Methodology",
    56: "4.1 Logit Margin Score (LMS)",
    57: "4.1.1 LMS Selection Rationale",
    58: "Table 3. LMS Selection Rationale",
    59: "4.1.2 Decision Token Filter",
    60: "**LMS** computation 이전에 tokens는 entropy-based criterion으로 filtering된다. Shannon entropy H(i) = −Σ pₖ log pₖ가 threshold θ_entropy를 초과하는 tokens만 포함한다. θ_entropy는 natural(**RAG**-off) responses의 entropy distribution에서 pre-registered 70th percentile로 **CR2** data에서 결정되며, **Run B** 이전에 고정되고 모든 conditions에서 유지된다. 이 percentile threshold는 near-deterministic top-1 probabilities를 갖는 high-frequency grammatical morpheme tokens를 억제하면서 semantic decision variance가 나타날 가능성이 큰 tokens를 보존하기 위해 사용된다.",
    61: "4.1.3 LMS Computation",
    62: "각 conversation turn t에 대해 **LMS**는 다음과 같이 계산된다.",
    64: "Table 4. LMS Primary Test Statistics",
    65: "**LMS**는 단독으로 causal evidence로 해석되지 않는다. **LMS** 증가는 additional retrieved information으로 인한 confidence 증가를 반영할 수 있다. 본 연구에서 **LMS**는 **Node C** 대비 differential pattern으로 나타나고, fixed policy reference를 향한 **CDS** movement와 co-occur하며, **CF-F**와 같은 counterfactual intervention conditions에서 replicated될 때에만 causally interpretable하다.",
    66: "4.2 Modality Analysis (MA)",
    67: "**Modality Analysis (MA)**는 Korean sentence-final endings를 기준으로 각 response sentence를 세 가지 epistemic categories 중 하나로 분류한다. measurement independence from the experimental variable를 유지하기 위해 **LLM**-based classification은 제외한다.",
    68: "Table 5. Modality Analysis Classification",
    69: "하나의 sentence가 여러 categories와 match되는 경우, 보수적 priority ordering인 Hedging > Epistemic > Assertive를 적용한다. lexical dictionary와 morphological analyzer version은 Git SHA-recorded되며 모든 runs에서 고정된다.",
    70: "4.2.1 MA Operational Definition",
    71: "각 **MA** variable은 turn별 classified sentences의 proportion으로 계산된다. S(t)를 turn t에서 response의 sentence set이라고 하고, |S_k(t)|를 type k로 분류된 sentences의 count라고 하자. operational formulas는 다음과 같다.",
    75: "구성상 각 sentence는 Hedging > Epistemic > Assertive priority rule 아래 정확히 하나의 category에 할당되므로, MA_assert(t) + MA_epist(t) + MA_hedge(t) = 1이다. measurement unit은 sentence이며 sentence boundaries는 morphological analyzer에 의해 결정된다.",
    76: "trigger threshold θ_MA는 **MA_assert**(t)에 적용되며, **Run B** 이전 **theta_config.json**에서 θ_LMS 및 θ_CDS와 동시에 고정된 natural(**RAG**-off) **Node C** responses의 **MA_assert** distribution 95th percentile로 **CR2** data에서 결정된다. post-CR2 modification은 허용되지 않는다.",
    77: "**RAG** injection 이후 **LMS_delta** increase와 동시에 발생하는 **MA_assert**(t) increase는 secondary causal evidence pattern(Section 4.6)을 구성하며, causal claims의 독립적 근거가 아니라 primary **LMS**/**CDS** pattern을 corroborate한다.",
    78: "4.3 Contextual Drift Score (CDS)",
    79: "**Contextual Drift Score (CDS)**는 각 turn response embedding과 reference policy objective embedding 사이의 cosine distance를 측정한다. reference embedding은 experimental runs 이전 ChromaDB corpus에 load된 모든 **RAG** document chunks의 mean embedding으로 구성되며, fixed embedding model(version-pinned and Git SHA-recorded)을 사용해 계산된다. 이 corpus-level mean embedding은 **CR** 이전 한 번 계산되고 모든 experimental conditions에서 고정된다. concurrent **Node C** baseline 대비 **CDS**가 statistically distinguishable하게 감소하고, 특히 retrieval-context injection과 temporally associated될 때, 이는 H1 아래 **context-level intervention**에 대한 multi-metric evidence pattern의 일부를 구성한다.",
    80: "4.4 Self-Reference Rate (SRR)",
    81: "**Self-Reference Rate (SRR)**는 semantically self-referential sentences의 proportion을 측정한다. 이는 preceding three turns(fixed window)의 어떤 sentence와도 cosine similarity > 0.85인 sentences로 정의된다. 추가로 N-gram(3) overlap criterion(> 0.4)을 만족하는 sentences는 simple repetition으로 분류되어 **SRR**에서 제외되며, 이를 통해 mechanical repetition과 의미 있는 self-reference를 분리한다. **SRR**은 secondary corroborating measure로 보고되며 causal claims를 독립적으로 지지하지 않는다.",
    82: "4.5 Structural Consistency Index (SCI)",
    83: "**Structural Consistency Index (SCI)**는 responses 안의 claim-evidence-conclusion transition pattern의 consistency를 측정한다. 이는 current turn의 transition matrix와 **CR2** 동안 **Node C** responses에서 도출한 baseline transition matrix 사이의 cosine similarity로 계산된다. main experimental runs 이전에 **SCI**는 pre-validity test(minimum 30 sentences, two independent annotators)를 거친다. inter-annotator agreement는 Cohen's κ로 평가하며, **SCI** inclusion에는 κ ≥ 0.6이 요구된다. 이 threshold가 충족되지 않으면 **SCI**는 pre-registered sentence-type proportion measure로 대체된다. 이 fallback procedure는 manual intervention 없이 deterministically triggered되어 post-hoc metric selection bias를 제거한다. **SCI**는 exploratory metric으로 취급되며 causal conclusions를 독립적으로 지지하지 않는다.",
    84: "4.6 Five-Dimensional Measurement Integration",
    85: "다섯 dimensions는 causal identification을 강화하기 위해 함께 해석된다. primary causal evidence pattern은 **LMS** increase alone이 아니라 **LMS_delta** increase, fixed policy reference를 향한 **CDS** decrease, 그리고 동일한 shift를 보이지 않는 matched **Node C** baseline의 joint observation이다. secondary evidence patterns에는 **LMS_delta** increase와 **MA_assert** increase의 concurrence, 그리고 intervention intensity와 일치하는 **Node A** > **Node B** > **Node C** ordering이 포함된다. **Node A** minus **Node B** contrast는 **RAG**-only intervention을 넘어서는 **SC-Protocol**의 differential contribution으로 해석되며, **CF-F** condition은 reverse causality에 대한 critical test를 제공한다. semantic intervention effect를 simple recency bias 또는 prompt-length distraction과 구분하기 위해, **Counterfactual** design은 content-level substitution controls(**CF-A**, **CF-C**, **CF-F**)를 critical baseline refutations로 포함한다. response structure shifts가 semantic content가 아니라 appended text의 존재만으로 발생한다면, content substitution과 opposing-document conditions는 directional differences를 산출하지 못할 것이다.",
    86: "5. Counterfactual Design",
    87: "5.1 Theoretical Foundation and Operational Intervention Definition",
    88: "causal identification strategy는 direct parameter-level intervention을 주장하기보다 observational conditioning과 controlled intervention을 구분하기 위한 theoretical reference point로 Pearl의 **do-calculus**를 사용한다. 본 연구에서 do(X)는 model parameters에 대한 intervention을 의미하지 않는다. 오히려 fixed model weights, fixed user utterances, deterministic decoding, matched concurrent baseline conditions 아래 retrieval-context provision을 externally controlled manipulation하는 것을 의미한다. 따라서 causal estimand는 model 자체를 바꾸는 효과가 아니라 externally supplied retrieval context가 observable response structure에 미치는 효과이다. **RAG**의 causal effect는 P(Y | do(X_context))가 otherwise fixed conditions 아래 matched baseline response distribution과 다를 때 operationally examined된다. underlying DAG는 다음과 같이 지정된다. U -> Y (direct path), U -> T -> X_context -> Y (mediated path), T <- {LMS, CDS} <- Y_{t-1} (feedback). 여기서 U는 model의 internal parametric knowledge state와 latent context accumulation을 포함한 unobserved confounders를 의미하고, T는 **LACP**/**SC-Protocol**이 통제하는 trigger assignment state를 나타내며, X_context는 do(X_context)로 operationalize되는 externally manipulated retrieval-context intervention이고, Y는 multi-dimensional observable response metrics(**LMS**, **MA**, **CDS**, **SRR**, **SCI**)를 의미한다. path T <- Y_{t-1} -> Y는 reverse causality analysis가 다루어야 하는 backdoor path를 구성한다. **CF-F**는 pre-selected non-trigger-eligible turns에서 retrieval-context injection을 강제함으로써 이 backdoor를 block하도록 특별히 설계되었다.",
    89: "5.2 Reverse Causality: Problem and Resolution",
    90: "\"**RAG**가 response structure를 바꾼 것인가, 아니면 response structure가 이미 바뀌고 있었고 그것이 trigger fire를 유발한 것인가?\"",
    91: "**CF-F**(Forced Random Intervention)는 experimental runs 이전에 randomly selected된 turns에서 **RAG**를 inject하며, seed-fixed assignments는 모든 **CF-F** repetitions에 일관되게 적용된다. injection turns는 **CR2** threshold calibration 동안 식별된 non-trigger-eligible positions에서 uniformly sampled된다. **CF-F**에서 significant **LMS_delta**가 관찰되면, 이는 reverse causality interpretation을 직접 반박한다.",
    92: "5.3 Six Counterfactual Conditions",
    93: "Table 6. Counterfactual Conditions",
    94: "6. 결과",
    95: "empirical results는 experimental runs 완료 후 보고될 예정이다. experimental schedule은 다음 순서로 진행된다. corpus hash fixation, PDF canonicalization and embedding pipeline verification, formal retrieval candidate freeze-gate verification, end-to-end 3-node system validation을 위한 **Test Run (TR)**(N=1-3); inter-node variance measurement를 위한 **Calibration Run (CR)**(N=10); natural metric distribution and threshold determination을 위한 **Calibration Run 2 (CR2)**(N=3); effect size estimation을 위한 **Run B** pilot(N=10); 그리고 power analysis로 sample size가 결정되는 main experimental runs. **Counterfactual** condition runs(**CF-A** through **CF-F**, N=5 each)는 main run completion 이후 수행된다.",
    96: "7. 논의",
    97: "discussion은 empirical result completion 이후 제공될 예정이다. discussion은 **RAG** system design에서 causal vs. augmentation distinction의 theoretical implications, on-premises public sector AI governance에 대한 practical implications, current experimental design의 limitations, 그리고 logit access가 unavailable한 cloud-based **LLM** environments로의 extension을 포함한 future research directions를 다룰 것이다.",
    98: "8. 결론",
    99: "본 논문은 on-premises **LLM** 기반 공공복지 상담 시스템에서 **Retrieval-Augmented Generation**이 externally controlled **context-level intervention**으로 기능하는지를 실증적으로 검토하기 위한 **LACP** experimental design을 제시했다. operationalized do-calculus-inspired framework에 기반한 proposed three-node simultaneous comparison architecture는 sequential comparison의 confounds 없이 controlled within-turn comparison을 가능하게 한다. entropy-based decision token filtering과 결합된 **LMS** metric은 Korean morphological token dilution을 다루지만, multi-metric causal evidence pattern의 일부로서만 해석된다. **CF-F** condition은 기존 **RAG** evaluation frameworks에 없던 reverse causality에 대한 direct test를 제공한다. Section 4.2.1은 **Modality Analysis**의 complete operational definition을 제공하며, **MA_assert**(t)를 **CR2** natural distribution data에서 도출된 pre-registered threshold **theta_MA**를 갖는 rule-based, **LLM**-independent proportion measure로 설정한다.",
    100: "감사의 글",
    101: "외부 연구비 지원은 없었다. 저자들은 manuscript organization, data-diffing assistance, experimental utilities를 위한 coding support, language polishing에 한정하여 AI-assisted coding and language tools를 사용했다. 모든 research questions, experimental design decisions, analytical interpretations, final manuscript responsibility는 저자들에게 있다.",
    102: "References",
}


TABLES = {
    0: [
        ["Node", "Role", "Intervention Level", "Trigger Monitoring"],
        ["Node A", "RAG + SC-Protocol", "Maximum", "Subject to monitoring"],
        ["Node B", "RAG only", "Intermediate", "Subject to monitoring"],
        ["Node C", "Baseline — Concurrent Control Observation", "None (fixed)", "Excluded"],
    ],
    1: [
        ["Phase", "Name", "Purpose", "N"],
        ["TR", "Test Run", "PDF canonicalization, chunking, embedding pipeline을 검증하고 corpus hash fixation(SHA-256 + Git SHA)을 확인하며, CR 이전 end-to-end 3-node system operation을 validate", "1-3"],
        ["CR", "Calibration Run", "inter-node hardware/runtime variance 측정", "10"],
        ["CR2", "Calibration Run 2", "natural metric variation 측정 → absolute thresholds(θ) 설정", "3"],
        ["Run B", "Main Experiment", "RAG pure effect와 SC-Protocol differential contribution을 동시에 측정", "10 (pilot)"],
        ["CF Runs", "Counterfactual Conditions", "6개 CF conditions(CF-A through CF-F)에 걸친 causal verification", "5 each"],
    ],
    2: [
        ["Node", "Initial", "Peak", "End before cooldown", "10-sec cooldown", "Drop"],
        ["inference1", "48.375 C", "74.0 C", "74.0 C", "67.5 C", "6.5 C"],
        ["inference2", "49.5 C", "69.0 C", "68.0 C", "64.0 C", "4.0 C"],
        ["inference3", "49.375 C", "74.0 C", "74.0 C", "67.375 C", "6.625 C"],
    ],
    3: [
        ["Dimension", "Justification"],
        ["Theoretical basis", "Pre-softmax logit values는 post-softmax probabilities보다 model의 internal confidence state를 더 직접적으로 반영한다 [13]. logit margin z1-z2는 saturation 없이 discrimination을 유지한다."],
        ["Korean morphological token dilution", "Korean-language outputs에는 near-deterministic top-1 probabilities를 갖는 grammatical morpheme tokens의 비율이 높다. entropy-based filtering을 적용한 LMS는 genuine probability mass redistribution이 발생하는 tokens를 분리한다."],
        ["Measurement objective alignment", "핵심 관심은 RAG intervention이 judgment-related tokens를 생성할 때 model의 confidence structure를 바꾸는지 여부이다. LMS는 response direction에 대한 model의 commitment를 직접 포착한다."],
    ],
    4: [
        ["Variable", "Formula", "Causal Interpretation"],
        ["LMS_delta_B(t)", "LMS_NodeB(t) − LMS_NodeC(t)", "concurrent baseline 대비 RAG-only differential effect"],
        ["LMS_delta_A(t)", "LMS_NodeA(t) − LMS_NodeC(t)", "concurrent baseline 대비 combined RAG + SC-Protocol differential effect"],
        ["LMS vs. Accuracy", "Pearson r(LMS, factual_accuracy)", "auxiliary evidence only: confidence-accuracy independence"],
        ["LMS_delta_AB(t)", "LMS_NodeA(t) - LMS_NodeB(t)", "secondary comparison: RAG-only intervention을 넘어서는 SC-Protocol의 differential contribution"],
    ],
    5: [
        ["Type", "Classification Criteria (Korean Sentence-Final Endings)", "Variable"],
        ["Assertive (단정)", "~입니다/습니다, ~해야 합니다/됩니다, ~불가합니다", "MA_assert(t)"],
        ["Epistemic (추측)", "~인 것 같습니다, ~으로 보입니다, ~추정됩니다", "MA_epist(t)"],
        ["Hedging (유보)", "~일 수 있습니다, 확인이 필요합니다", "MA_hedge(t)"],
    ],
    6: [
        ["CF Condition", "Design", "Causal Contribution", "N"],
        ["CF-A: Content Substitution", "same context, different RAG document(C↔D welfare program pair)", "direct verification: content change → response change", "5"],
        ["CF-B: Empty Result", "same context, null retrieval result injected", "baseline: RAG-absent counterfactual comparison", "5"],
        ["CF-C: Opposing Document", "previous-version policy document injected(date identifiers removed)", "direction reversal test — H1에 대한 가장 강한 evidence", "5"],
        ["CF-D: Temporal Shift", "trigger와 무관하게 turns 5, 15, 25에서 RAG injection forced", "injection timing effect + partial reverse causality resolution", "5"],
        ["CF-E: Internal-External Separation", "user queries의 internal contextual cues는 고정하고, external retrieval payloads만 substitute하여 internal reasoning과 external intervention을 분리", "external attribution confirmation — causal foundation", "5"],
        ["CF-F: Forced Random Intervention", "uniformly sampled non-trigger-eligible turns에서 RAG injected(seed-fixed, CR2-verified)", "direct reverse causality refutation — Primary Test 2", "5"],
    ],
}


def split_marked(text):
    parts = []
    i = 0
    while i < len(text):
        start = text.find("**", i)
        if start == -1:
            parts.append((text[i:], False))
            break
        if start > i:
            parts.append((text[i:start], False))
        end = text.find("**", start + 2)
        if end == -1:
            parts.append((text[start:], False))
            break
        parts.append((text[start + 2:end], True))
        i = end + 2
    return [(p, b) for p, b in parts if p]


def apply_east_asia_font(run, font_name="Malgun Gothic"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_para_text(paragraph, text):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    for part, bold in split_marked(text):
        run = paragraph.add_run(part)
        run.bold = bold
        apply_east_asia_font(run)


def bold_terms_in_paragraph(paragraph):
    text = paragraph.text
    if not text:
        return
    marked = text
    for term in sorted(KEY_TERMS, key=len, reverse=True):
        marked = marked.replace(term, f"**{term}**")
    if marked != text:
        set_para_text(paragraph, marked)


def main():
    shutil.copyfile(SRC, OUT)
    doc = docx.Document(OUT)

    for idx, text in PARA.items():
        set_para_text(doc.paragraphs[idx], text)

    for ti, rows in TABLES.items():
        table = doc.tables[ti]
        for r, row in enumerate(rows):
            for c, text in enumerate(row):
                cell = table.cell(r, c)
                para = cell.paragraphs[0]
                set_para_text(para, text)
                for extra in cell.paragraphs[1:]:
                    set_para_text(extra, "")
                bold_terms_in_paragraph(para)

    # Reinforce bolding for any key terms in translated body before References only.
    for p in doc.paragraphs[:103]:
        if "**" not in p.text:
            bold_terms_in_paragraph(p)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
