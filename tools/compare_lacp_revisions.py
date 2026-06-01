from __future__ import annotations

import argparse
import difflib
import json
import re
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


KEYWORDS = [
    "Pre-CR",
    "Test Run",
    "TR",
    "Calibration Run",
    "CR2",
    "Run B",
    "Node C",
    "threshold",
    "theta",
    "LMS",
    "CDS",
    "MA",
    "counterfactual",
    "top-k",
    "retrieval",
    "hybrid",
    "fallback",
    "corpus",
    "hash",
    "chunk",
    "embedding",
    "context window",
    "truncation",
    "done_reason",
    "deterministic",
    "clock",
    "NTP",
    "Vulkan",
    "Ollama",
    "MariaDB",
    "ChromaDB",
    "RAG-off",
    "RAG",
]


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def iter_docx_blocks(path: Path) -> list[str]:
    if not zipfile.is_zipfile(path):
        text = path.read_text(encoding="utf-8", errors="replace")
        return [clean_text(line) for line in text.splitlines() if clean_text(line)]
    doc = Document(str(path))
    blocks: list[str] = []
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            blocks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                blocks.append(row_text)
    return blocks


def split_sentences(blocks: list[str]) -> list[str]:
    sentences: list[str] = []
    for block in blocks:
        parts = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9(])", block)
        for part in parts:
            part = clean_text(part)
            if len(part) > 25:
                sentences.append(part)
    return sentences


def likely_headings(blocks: list[str]) -> list[str]:
    headings: list[str] = []
    for block in blocks:
        if re.match(r"^(\d+(\.\d+)*\.?\s+|Abstract$|References$|Keywords\b)", block):
            headings.append(block)
        elif len(block) < 95 and re.search(r"(Method|Measurement|Protocol|Result|Discussion|Validation|Threshold|Run|Analysis|Limitations)", block):
            headings.append(block)
    return headings


def keyword_counts(text: str) -> dict[str, int]:
    counts: dict[str, int] = {}
    lower = text.lower()
    for keyword in KEYWORDS:
        counts[keyword] = lower.count(keyword.lower())
    return counts


def meaningful_added_removed(prev: list[str], cur: list[str], limit: int = 22) -> tuple[list[str], list[str]]:
    prev_set = set(prev)
    cur_set = set(cur)
    added = [s for s in cur if s not in prev_set]
    removed = [s for s in prev if s not in cur_set]

    def score(sentence: str) -> tuple[int, int]:
        keyword_hit = sum(1 for kw in KEYWORDS if kw.lower() in sentence.lower())
        return keyword_hit, min(len(sentence), 500)

    added = sorted(added, key=score, reverse=True)[:limit]
    removed = sorted(removed, key=score, reverse=True)[:limit]
    return added, removed


def version_label(path: Path) -> str:
    name = path.stem.replace("lacp_ijibc_", "")
    return name


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--folder", required=True)
    parser.add_argument("--out", required=True)
    parser.add_argument("files", nargs="+")
    args = parser.parse_args()

    folder = Path(args.folder)
    paths = [folder / file_name for file_name in args.files]

    docs = []
    for path in paths:
        blocks = iter_docx_blocks(path)
        sentences = split_sentences(blocks)
        full_text = "\n".join(blocks)
        docs.append(
            {
                "file": str(path),
                "label": version_label(path),
                "word_count": len(re.findall(r"\b[\w-]+\b", full_text)),
                "paragraph_count": len(blocks),
                "headings": likely_headings(blocks),
                "keywords": keyword_counts(full_text),
                "sentences": sentences,
            }
        )

    comparisons = []
    for prev, cur in zip(docs, docs[1:]):
        added, removed = meaningful_added_removed(prev["sentences"], cur["sentences"])
        prev_keys = prev["keywords"]
        cur_keys = cur["keywords"]
        keyword_delta = {
            key: cur_keys[key] - prev_keys[key]
            for key in KEYWORDS
            if cur_keys[key] - prev_keys[key] != 0
        }
        matcher = difflib.SequenceMatcher(a=prev["sentences"], b=cur["sentences"])
        comparisons.append(
            {
                "from": prev["label"],
                "to": cur["label"],
                "word_delta": cur["word_count"] - prev["word_count"],
                "similarity": round(matcher.ratio(), 3),
                "keyword_delta": dict(sorted(keyword_delta.items(), key=lambda item: abs(item[1]), reverse=True)),
                "added": added,
                "removed": removed,
                "new_headings": [h for h in cur["headings"] if h not in set(prev["headings"])],
            }
        )

    first = docs[0]
    last = docs[-1]
    total_keyword_delta = {
        key: last["keywords"][key] - first["keywords"][key]
        for key in KEYWORDS
        if last["keywords"][key] - first["keywords"][key] != 0
    }

    report = {
        "docs": [{k: v for k, v in doc.items() if k != "sentences"} for doc in docs],
        "comparisons": comparisons,
        "overall": {
            "from": first["label"],
            "to": last["label"],
            "word_delta": last["word_count"] - first["word_count"],
            "keyword_delta": dict(sorted(total_keyword_delta.items(), key=lambda item: abs(item[1]), reverse=True)),
            "new_headings": [h for h in last["headings"] if h not in set(first["headings"])],
        },
    }

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
