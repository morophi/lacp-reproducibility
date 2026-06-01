from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\morophi\OneDrive\문서\New project\LACP_Professor_Meeting_Review_Plan_2026-05-30.docx")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = False
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell(cell, h, True)
        shade(cell, "F2F4F7")
        cell.width = Inches(widths[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
            cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style in ["List Bullet", "List Number"]:
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_callout(doc: Document, title: str, body: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    shade(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 58, 95)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    doc.add_paragraph()


def footer(doc: Document) -> None:
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("LACP professor meeting review plan")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(85, 85, 85)


def build() -> None:
    doc = Document()
    styles(doc)
    footer(doc)

    doc.add_paragraph("LACP Professor Meeting Review Plan", style="Title")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Meeting date: May 30, 2026 | Purpose: manuscript, experiment readiness, and file-version recovery review").italic = True

    add_callout(
        doc,
        "Meeting goal",
        "내일 교수님과의 검토는 단순히 최신본을 다시 설명하는 자리가 아니라, 최신 원고의 제출 가능성, 실험 근거의 방어 가능성, 그리고 산재한 버전/증거 문서의 정리 방향을 확정하는 자리로 잡는다.",
    )

    doc.add_heading("1. 교수님께 먼저 말씀드릴 핵심 상황", level=1)
    bullets(
        doc,
        [
            "최신본은 전송했지만, rev2부터 rev9까지 수정 흐름이 길어지면서 어떤 변경이 본문에 반영되었고 어떤 내용이 evidence 문서에만 남아 있는지 경계가 흐려졌다.",
            "특히 logprobs 수집 오류와 해결책은 revision note 초안에는 빠져 있었고, 실제 근거는 E2E/Smoke evidence 문서에 더 명확히 남아 있다.",
            "따라서 교수님께는 최신 원고 자체뿐 아니라, 어떤 operational readiness evidence를 함께 제시해야 하는지 확인받는 것이 필요하다.",
            "현재 목표는 논문 내용을 더 늘리는 것이 아니라, 원고-실험가이드-체크리스트-evidence 문서 사이의 대응관계를 회복하는 것이다.",
        ],
    )

    doc.add_heading("2. 회의에서 반드시 확인받을 결정 사항", level=1)
    table(
        doc,
        ["우선순위", "확인할 질문", "교수님께 받고 싶은 결정"],
        [
            [
                "1",
                "rev9 원고를 최신 제출본으로 고정해도 되는가?",
                "최신 manuscript 기준 파일을 lacp_ijibc_rev9.docx로 확정할지, 또는 rev9_kor_merge/별도 교정본을 참고해야 할지 결정.",
            ],
            [
                "2",
                "revision note에 어느 수준까지 기술적 시행착오를 넣을 것인가?",
                "logprobs endpoint correction, CDS reference artifact, top-k feasibility, thermal safeguard를 note에 포함할 범위 확정.",
            ],
            [
                "3",
                "logprobs 오류/해결책을 논문 본문에 반영해야 하는가, 아니면 supplement/evidence note로 충분한가?",
                "본문에는 measurement endpoint policy로 간단히 넣고, 상세 오류는 evidence note로 분리할지 결정.",
            ],
            [
                "4",
                "TR/CR/CR2/Run B 실행 전 문서 freeze 기준은 무엇인가?",
                "원고, experiment guideline, node checklist, harness evidence의 freeze 순서와 naming rule 확정.",
            ],
            [
                "5",
                "교수님께 추가로 전달할 패키지는 무엇인가?",
                "최신 원고, revision note, meeting review plan, E2E issue summary, top-k=3 smoke evidence 중 전달 파일 묶음 확정.",
            ],
        ],
        [0.75, 2.75, 3.0],
    )

    doc.add_heading("3. 기술적으로 꼭 짚어야 할 쟁점", level=1)
    table(
        doc,
        ["쟁점", "현재 확인된 내용", "회의에서의 처리 방향"],
        [
            [
                "LMS/logprobs",
                "native Ollama /api/generate 또는 native chat 경로는 token-level logprobs/top_logprobs를 제공하지 않아 LMS 입력이 비었다. 해결책은 OpenAI-compatible /v1/chat/completions, logprobs=true, top_logprobs=5, think=false 사용.",
                "본문에 endpoint capability control로 짧게 반영할지, evidence note에만 둘지 결정.",
            ],
            [
                "thinking artifact",
                "Qwen 계열에서 empty <think></think> shell이 나타날 수 있어 response text와 LMS token position을 오염시킬 위험이 있었다.",
                "empty shell strip, 해당 prefix token LMS 제외, non-empty thinking은 failed TR 조건으로 유지.",
            ],
            [
                "top-k",
                "top_k=5는 현재 BC-250/Ollama/Qwen3 환경에서 안정성 문제가 있었고, top_k=3 smoke pass가 metric-complete evidence로 남아 있다.",
                "formal TR 후보 top_k를 3으로 둘지, 5를 보류 조건으로 유지할지 확인.",
            ],
            [
                "CDS reference",
                "CDS reference embedding artifact가 없으면 CDS null이 발생한다. frozen corpus stored embeddings 평균으로 reference artifact를 만들고 Harness는 read-only로 사용해야 한다.",
                "corpus/reference artifact freeze 절차를 guideline/checklist에 명시.",
            ],
            [
                "thermal safeguard",
                "parallel A/B/C generation, logprobs, RAG prompts는 온도 상승과 node stability에 영향을 준다. cooldown policy는 실험 처치가 아니라 operational control로 기록되어야 한다.",
                "rev8.1/updated note의 thermal safeguard 설명을 교수님께 확인.",
            ],
            [
                "threshold freezing",
                "rev9에서 theta_entropy와 theta_LMS/CDS/MA를 분리했고, trigger/effect threshold는 Node C-relative oriented differential score로 CR2에서 고정한다.",
                "rev9 Section 4.6이 현재 방법론 방어의 핵심임을 설명하고 승인받기.",
            ],
        ],
        [1.35, 3.35, 1.8],
    )

    doc.add_heading("4. 버전관리 복구 계획", level=1)
    add_callout(
        doc,
        "현재 문제 인식",
        "파일명을 나름 관리했지만 rev7.x, rev8.1 variants, Korean merge, evidence note, experiment guideline, node checklist가 동시에 늘어나면서 무엇이 manuscript 변경이고 무엇이 실행/검증 evidence인지 구분이 약해졌다.",
    )
    numbered(
        doc,
        [
            "Manuscript line을 하나로 고정한다: lacp_ijibc_rev9.docx를 기준 후보로 두고, 변형본은 translation/merge/reference로만 표시한다.",
            "Evidence line을 별도 고정한다: E2E_ISSUE_SAMPLING_SUMMARY_20260527.md, SMOKE_EVIDENCE_TOPK3_20260527.md, harness_file_structure_and_db_flow_rev3.1.md를 operational evidence 묶음으로 둔다.",
            "Execution guide line을 분리한다: lacp_experiment_guideline_rev5.6.md와 lacp_node_checklist_v10.5.md는 실행 절차 문서로 두고, logprobs 오류의 상세 원인은 evidence note에서 참조한다.",
            "교수님 확인 후 freeze folder를 만든다: manuscript, revision_note, evidence, execution_guide 네 폴더로 분리하고 파일명에 날짜와 역할을 붙인다.",
            "다음부터는 파일명에 역할을 붙인다: manuscript_rev, revision_note, evidence_smoke, issue_summary, guideline, checklist를 섞지 않는다.",
        ],
    )

    doc.add_heading("5. 교수님께 보여드릴 파일 묶음", level=1)
    table(
        doc,
        ["파일", "역할", "회의에서 확인할 점"],
        [
            ["lacp_ijibc_rev9.docx", "최신 영문 원고 후보", "이 파일을 최신본으로 고정해도 되는지 확인."],
            ["LACP_IJIBC_Revision_Note_rev2_to_rev9_updated.docx", "rev2-rev9 변경 요약", "logprobs 보완 섹션 포함 수준이 적절한지 확인."],
            ["E2E_ISSUE_SAMPLING_SUMMARY_20260527.md", "오류/해결책 evidence", "native endpoint logprobs 부재와 OpenAI-compatible endpoint 전환 근거 설명."],
            ["SMOKE_EVIDENCE_TOPK3_20260527.md", "top-k=3 metric-complete smoke pass", "formal TR 전 operational readiness evidence로 제시 가능 여부 확인."],
            ["lacp_experiment_guideline_rev5.6.md", "실험 실행 가이드", "본문/체크리스트/evidence와 충돌하거나 빠진 항목 확인."],
            ["lacp_node_checklist_v10.5.md", "노드 실행 체크리스트", "logprobs endpoint policy가 충분히 반영되어 있는지 확인."],
        ],
        [2.25, 2.0, 2.25],
    )

    doc.add_heading("6. 회의 진행 순서 제안", level=1)
    numbered(
        doc,
        [
            "먼저 최신본 기준을 확정한다: rev9를 기준으로 볼지, Korean merge 또는 별도 교정본을 참고할지 확인.",
            "Revision note의 목적을 설명한다: 단순 변경 내역이 아니라 methodological defensibility를 보여주는 보조 문서라고 설명.",
            "logprobs 문제를 짧게 보고한다: native endpoint가 LMS 입력을 제공하지 않았고, formal endpoint 전환으로 metric-complete smoke pass가 가능해졌다고 설명.",
            "교수님께 논문 본문 반영 범위를 묻는다: endpoint correction을 방법론 본문에 넣을지, supplementary/evidence note로 둘지 결정.",
            "마지막으로 버전관리 복구 계획을 제안한다: manuscript/evidence/guideline/checklist를 분리해 freeze package로 재정리.",
        ],
    )

    doc.add_heading("7. 교수님께 드릴 수 있는 짧은 설명", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "교수님, 최신본을 보내드린 뒤 다시 정리해 보니 원고 수정본과 실험 evidence 문서가 일부 섞여 있어서, 내일은 최신 원고의 방향뿐 아니라 어떤 evidence를 함께 제시해야 하는지도 확인받고 싶습니다. 특히 LMS 계산에 필요한 logprobs 수집은 native Ollama endpoint에서 막혔고, OpenAI-compatible chat completions 경로로 전환하면서 metric-complete smoke evidence를 확보했습니다. 이 내용을 본문에 어느 정도 반영할지, 또는 별도 readiness note로 둘지 교수님 의견을 듣고 싶습니다."
    )

    doc.add_heading("8. 회의 후 바로 할 일", level=1)
    table(
        doc,
        ["Action", "Output", "Timing"],
        [
            ["교수님 결정사항 기록", "meeting_decisions_2026-05-30.md", "회의 직후"],
            ["최신 원고 기준 파일명 고정", "lacp_ijibc_manuscript_freeze_20260530.docx", "회의 당일"],
            ["Revision note 최종화", "LACP_IJIBC_Revision_Note_final_20260530.docx", "회의 당일"],
            ["Evidence package 정리", "evidence/ 폴더와 README", "회의 후 1일 이내"],
            ["Guideline/checklist gap 반영", "experiment guideline + node checklist patch note", "회의 후 1-2일 이내"],
        ],
        [2.25, 2.75, 1.5],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))


if __name__ == "__main__":
    build()
