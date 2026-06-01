from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\morophi\OneDrive\문서\New project\LACP_IJIBC_Revision_Note_rev2_to_rev9_updated.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, True)
        set_cell_shading(cell, "E8EEF5")
        cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(10)

    for name, size, before, after in [
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("LACP IJIBC Revision Note")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(85, 85, 85)


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Revision Note for LACP IJIBC Manuscript")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Scope: comparative review of lacp_ijibc_rev2 through lacp_ijibc_rev9").italic = True
    subtitle.paragraph_format.space_after = Pt(12)

    meta = doc.add_paragraph()
    meta.add_run("Prepared for: Professor\n").bold = True
    meta.add_run(f"Prepared date: {date(2026, 5, 29).strftime('%B %d, %Y')}\n")
    meta.add_run("Latest manuscript checked: lacp_ijibc_rev9.docx")

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This revision note summarizes the substantive changes made across the LACP IJIBC manuscript from "
        "rev2 to rev9. The note is intended to accompany the latest manuscript version and clarify that the "
        "recent revisions were not cosmetic only; they strengthened the paper's causal identification logic, "
        "operational validity controls, measurement definitions, and reproducibility safeguards."
    )

    doc.add_heading("Scope of Comparison", level=1)
    add_bullets(
        doc,
        [
            "Compared the main manuscript line from lacp_ijibc_rev2.docx to lacp_ijibc_rev9.docx, including major intermediate revisions rev3, rev4, rev5, rev5.1/rev5_1, rev6, rev7, rev7.1 to rev7.9, rev8.1, and rev9.",
            "Focused on substantive manuscript evolution: causal framing, experimental architecture, pre-run validation, measurement formulas, threshold definitions, retrieval controls, hardware/runtime safeguards, and section structure.",
            "Excluded Korean merge/translation variants from the main comparison path, except as contextual evidence, because the latest English manuscript line culminates in lacp_ijibc_rev9.docx.",
        ],
    )

    doc.add_heading("Executive Summary", level=1)
    add_bullets(
        doc,
        [
            "The manuscript evolved from a protocol-centered experimental draft into a more defensible causal-intervention study using Node C as a concurrent no-intervention baseline.",
            "The experimental procedure now includes a clearly defined Pre-CR Test Run (TR) that verifies corpus hash fixation, PDF canonicalization, chunking, embedding, ChromaDB ingestion, Harness orchestration, metric logging, and MariaDB writes before formal calibration begins.",
            "The measurement layer was strengthened by expanding LMS, MA, CDS, SRR, and SCI definitions, with particular emphasis on MA as a rule-based, LLM-independent modality proportion measure.",
            "The latest rev9 closes the threshold-definition gap by distinguishing theta_entropy from intervention-effect thresholds and by defining theta_LMS, theta_CDS, and theta_MA as CR2-frozen Node C-relative differential thresholds.",
            "Operational validity controls were added for reasoning-output contamination, clock synchronization, top-k payload feasibility, context-window saturation, hybrid table-exposure fallback logging, truncation monitoring, and thermal stability.",
        ],
    )

    doc.add_heading("Revision Trajectory", level=1)
    add_table(
        doc,
        ["Revision", "Main revision focus", "Reason this matters"],
        [
            [
                "rev2 to rev4",
                "Strengthened AI-governance and SC-Protocol framing; clarified that SC-Protocol functions as an intervention-governing constraint layer rather than a broad protocol claim.",
                "Improves theoretical positioning and narrows the manuscript's claims to the empirical study.",
            ],
            [
                "rev5 to rev5.1",
                "Expanded do-calculus framing, injection sampling rules, reference embedding construction, and hardware/runtime details including Ollama/Vulkan device binding.",
                "Moves the draft toward a reproducible causal experiment rather than a descriptive RAG implementation.",
            ],
            [
                "rev6",
                "Consolidated manuscript into IJIBC-style structure and added complete MA operational definition, threshold sentence, and SCI measurement logic.",
                "Improves journal readability and makes the measurement layer auditable.",
            ],
            [
                "rev7 to rev7.3",
                "Introduced TR validation, corpus hash fixation, chunking/embedding audit, pipeline integrity checks, reasoning-output suppression, clock synchronization, top-k feasibility, context-window failure handling, hybrid fallback logging, and truncation monitoring.",
                "Prevents calibration and Run B from proceeding on an unstable or contaminated retrieval/inference pipeline.",
            ],
            [
                "rev7.4 to rev7.9",
                "Tightened operational do-calculus language, added entropy-based LMS token eligibility, strengthened SCM interpretation, added retrieval freeze-gate review, and integrated additional RAG/logit-confidence literature.",
                "Improves causal interpretation, literature grounding, and pre-registration discipline.",
            ],
            [
                "rev8.1",
                "Added thermal safeguard and turn-level synchronization section, including measured nonlinear post-load cooling behavior and uniform cooldown policy.",
                "Documents hardware stability as an operational validity control rather than an uncontrolled runtime nuisance.",
            ],
            [
                "E2E smoke evidence, 2026-05-27",
                "Identified that native Ollama `/api/generate`/native chat output did not expose token-level logprobs required for LMS; Harness was switched to OpenAI-compatible `/v1/chat/completions` with `logprobs=true`, `top_logprobs=5`, and `think=false`.",
                "Turns a previously blocked LMS/metric-completeness condition into an explicit endpoint-capability control before formal TR/CR execution.",
            ],
            [
                "rev9",
                "Added Section 4.6 Threshold Estimation and Freezing; separated theta_entropy from theta_LMS/theta_CDS/theta_MA; defined Node C-relative oriented differential scores and theta_config.json freezing.",
                "Closes a key methodological gap by aligning thresholds with the paper's causal contrast.",
            ],
        ],
        [1.15, 3.05, 2.25],
    )

    doc.add_heading("Major Substantive Revisions", level=1)
    doc.add_heading("1. Causal Design and Baseline Interpretation", level=2)
    doc.add_paragraph(
        "The revision line increasingly clarifies that the LACP experiment is not testing raw metric increases in isolation. "
        "Instead, RAG effects are interpreted against the matched concurrent Node C baseline under fixed user utterances, fixed model weights, deterministic decoding, and synchronized turn-level execution. "
        "This strengthens the causal logic by making the no-intervention baseline explicit throughout the threshold and measurement architecture."
    )

    doc.add_heading("2. Pre-Calibration Operational Validation", level=2)
    doc.add_paragraph(
        "A major addition beginning in rev7 is the Pre-CR Test Run (TR). TR verifies corpus integrity, pipeline execution, retrieval behavior, and logging completeness before CR, CR2, Run B, or counterfactual runs are accepted. "
        "The purpose is to separate infrastructure readiness from empirical causal evidence and avoid treating failed or partially logged runs as valid observations."
    )

    doc.add_heading("3. Retrieval and Corpus Controls", level=2)
    add_bullets(
        doc,
        [
            "Corpus hash fixation and version-pinned embedding model records were added before calibration.",
            "Top-k acceptance now requires not only retrieval completeness but also context-window feasibility, non-null generation, and complete metric-row production.",
            "Context-window saturation and prompt truncation are classified as retrieval substrate failures, not as evidence that RAG failed conceptually.",
            "Hybrid table-exposure fallback is allowed only as a minimal TR-stage exposure check, with vector-only results preserved separately and fallback use explicitly logged.",
            "Formal retrieval freeze-gate review was added to confirm section-, chunk-, block-type-, and table-level evidence exposure.",
        ],
    )

    doc.add_heading("4. Measurement Layer Refinement", level=2)
    doc.add_paragraph(
        "The measurement layer was expanded and clarified across revisions. LMS, MA, CDS, SRR, and SCI were retained as the five-dimensional measurement architecture, but later versions provide stronger operational definitions and safeguards. "
        "In particular, MA was clarified as a rule-based, LLM-independent proportion measure rather than a subjective modality judgment. CDS was tied more explicitly to movement toward a fixed policy reference relative to Node C, and SCI received fallback logic where transition-pattern thresholds are not satisfied."
    )

    doc.add_heading("4.1 LMS Logprobs Collection Error and Endpoint Correction", level=3)
    doc.add_paragraph(
        "The E2E readiness review identified a specific LMS collection failure that was not captured in the first version of this revision note. "
        "Earlier smoke rows using the native Ollama chat/generate path produced generated text but did not expose the token-level candidate distributions required for LMS. "
        "The observed symptom was zero-length raw and cleaned logprob arrays, resulting in missing LMS inputs and incomplete metric rows."
    )
    add_table(
        doc,
        ["Issue observed", "Cause", "Resolution / policy"],
        [
            [
                "Native Ollama endpoint returned response/context fields but not usable token-level `logprobs` or `top_logprobs`; validators showed metric-complete rows blocked.",
                "LMS requires per-token top candidate logprobs, not only generated text or token ids. The native endpoint did not satisfy that measurement requirement.",
                "Harness formal runtime was switched to OpenAI-compatible `/v1/chat/completions`; requests now use `logprobs=true`, `top_logprobs=5`, and `think=false`, with LMS extracted from `choices[0].logprobs.content`.",
            ],
            [
                "Empty `<think></think>` shell could enter response text or token positions even when substantive thinking content was disabled.",
                "Qwen template behavior can emit an empty thinking shell as a formatting artifact.",
                "Empty shells are stripped, matching prefix token positions are excluded from LMS, and non-empty thinking content remains a failed TR/quality condition.",
            ],
        ],
        [2.05, 2.0, 2.4],
    )
    doc.add_paragraph(
        "This correction should be described as an operational readiness safeguard rather than as a substantive causal result. "
        "It shows that the experiment now has a defined LMS-capable endpoint path and an auditable logprob extraction policy before formal runs begin."
    )

    doc.add_heading("5. Threshold Estimation and Freezing", level=2)
    doc.add_paragraph(
        "The most important final-stage revision is rev9's new Section 4.6. It distinguishes theta_entropy, which is only an LMS-internal token-eligibility filter estimated from CR2 natural RAG-off token entropy, from theta_LMS, theta_CDS, and theta_MA, which are trigger and intervention-effect thresholds. "
        "The latter thresholds are now estimated as empirical 95th percentiles over absolute Node C-relative differential score distributions and frozen before Run B."
    )
    add_table(
        doc,
        ["Metric", "Oriented differential score added in rev9", "Interpretation"],
        [
            ["LMS", "D_LMS^X(t) = LMS_X(t) - LMS_C(t)", "Larger value indicates stronger LMS commitment relative to Node C."],
            ["CDS", "D_CDS^X(t) = CDS_C(t) - CDS_X(t)", "Larger value indicates stronger movement toward the fixed policy reference relative to Node C."],
            ["MA", "D_MA^X(t) = MA_assert_X(t) - MA_assert_C(t)", "Larger value indicates stronger assertive modality relative to Node C."],
        ],
        [0.9, 2.55, 3.0],
    )

    doc.add_heading("6. Runtime and Hardware Validity Controls", level=2)
    doc.add_paragraph(
        "Later revisions also document runtime controls that are essential for a local multi-node LLM experiment. These include reasoning/thinking-output suppression, chrony/NTP timestamp interpretation, deterministic decoding, node output completeness checks, and rev8.1's thermal safeguard. "
        "The thermal control is important because post-load cooling behavior is treated as a fixed operational condition across TR, CR, CR2, CF, and CF-F runs rather than as an experimental treatment."
    )

    doc.add_heading("Recommended Cover Message", level=1)
    doc.add_paragraph(
        "Professor, I am sending this short revision note to accompany the latest version of the manuscript. "
        "The revisions from rev2 to rev9 mainly strengthened the methodological defensibility of the paper: the latest version clarifies the Node C-relative causal contrast, adds pre-calibration operational validation, makes retrieval and runtime safeguards explicit, and freezes all trigger/effect thresholds using the CR2-derived oriented-score convention. "
        "I believe these changes make the current version more suitable for review because they address reproducibility, measurement validity, and possible implementation-level confounds in a more explicit way."
    )

    doc.add_heading("Files Compared", level=1)
    add_bullets(
        doc,
        [
            "lacp_ijibc_rev2.docx, rev3.docx, rev4.docx, rev5.docx, rev5.1/rev5_1.docx, rev6.docx",
            "lacp_ijibc_rev7.docx, rev7.1.docx, rev7.2.docx, rev7.3.docx, rev7.4.docx, rev7.5.docx, rev7.7.docx, rev7.8.docx, rev7.9_5th_revise.docx",
            "lacp_ijibc_rev8.1.docx and lacp_ijibc_rev9.docx",
            "E2E_ISSUE_SAMPLING_SUMMARY_20260527.md, SMOKE_EVIDENCE_TOPK3_20260527.md, and harness_file_structure_and_db_flow_rev3.1.md for the logprobs endpoint correction and smoke-readiness evidence.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))


if __name__ == "__main__":
    build_doc()
